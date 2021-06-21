import urllib
import urllib.request
import re
import json
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import RGBColor,Pt
from urllib.request import urlopen
class Tool:
    '''
    用于数据清洗，删除无关符号
    '''
     # 去除img标签,7位长空格
    removeImg = re.compile('<img.*?>| {7}|')
    # 删除超链接标签
    removeAddr = re.compile('<a.*?>|</a>')
    # 把换行的标签换为\n
    replaceLine = re.compile('<tr>|<div>|</div>|</p>')
    # 将表格制表<td>替换为\t
    replaceTD = re.compile('<td>')
    # 把段落开头换为\n加空两格
    replacePara = re.compile('<p.*?>')
    # 将换行符或双换行符替换为\n
    replaceBR = re.compile('<br><br>|<br>')
    # 将其余标签剔除
    removeExtraTag = re.compile('<.*?>')

    def replace(self, x):
        x = re.sub(self.removeImg, "", x)
        x = re.sub(self.removeAddr, "", x)
        x = re.sub(self.replaceLine, "\n", x)
        x = re.sub(self.replaceTD, "\t", x)
        x = re.sub(self.replacePara, "\n    ", x)
        x = re.sub(self.replaceBR, "\n", x)
        x = re.sub(self.removeExtraTag, "", x)
        # strip()将前后多余内容删除
        return x.strip()

class XLBKCrawler:
    def __init__(self, baseUrl, articleTag, fileName):
        self.baseURL = baseUrl
        self.tool = Tool()
        self.file = None
        self.article = 1
        self.defaultTitle = u'新浪博客'
        self.articleTag = articleTag
        self.fileName = fileName

    def getPage(self, pageNum):
        '''
        获得url下网站最原始网页信息
        '''
        try:
            url = self.baseURL+str(pageNum)+'.html'
            # 即将抓取的网页
            # http://blog.sina.com.cn/s/articlelist_1215172700_10_1.html
            # 向网站发送访问请求
            request = urllib.request.Request(url)
            # 网站对我们访问请求的返回结果
            response = urllib.request.urlopen(request)
            return response.read().decode('utf-8')

        except urllib.error.URLError as e:
            if hasattr(e, "reason"):
                print("连接新浪博客失败,错误原因", e.reason)
                return None

    def getPageNum(self, page):
        '''
        page 是最原始的HTML网页代码
        返回值是当前的文章页码
        '''
        pattern = re.compile(u'<span style.*?>共(.*?)页</span>', re.S)
        result = re.search(pattern, page)
        if result:
            return result.group(1).strip()
        else:
            print(result)
            return 1
    def getUrl(self, page):
        '''
        获得所有文章的网页链接
        '''
        pattern = re.compile(
            '<span class="atc_title">.*?<a.*?href="(.*?)">.*?</a>.*?</span>', re.S)
        items = re.findall(pattern, page)
        urls = []
        for item in items:
            url = item
            urls.append(url)
        return urls
    def getText(self, url):
         '''
         获得文章正文
         '''
         text = urlopen(url).read().decode('utf-8')
         start = text.find(u"<!-- 正文开始 -->")
         end = text.find(u"<!-- 正文结束 -->")
         text = text[start:end]
         text = re.sub(re.compile('<p.*?>'), "\n    ", text)
         text = re.sub(re.compile('<p>'), "\n    ", text)
         text = re.sub(r'<(S*?)[^>]*>.*?|<.*? /> ', '', text)
         text = re.sub(r'&[^>]*?\;', ' ', text)
         text = re.sub('浏览“缠中说禅”更多文章请点击进入缠中说禅','',text)
        #  text = re.sub(r'[\n,\t]{2,}', '\n', text)
         text_split = [x for x in re.split(r'\s',text) if x]
         return text_split
    def getTime(self,url):
        text=urlopen(url).read().decode('utf-8')
        # <span class="time SG_txtc">(2008-10-10 09:24:36)</span><div class="turnBoxzz">
        # 提取时间
        pattern_time =re.compile('<span class="time SG_txtc">\((.*?)\)</span><div class="turnBoxzz">',re.S)
        time = ''.join(re.findall(pattern_time,text))
        return time
  
    def getTitle(self,url):
        text=urlopen(url).read().decode('utf-8')
        # <h2 id="t_486e105c0100aays" class="titName SG_txta">严重预告:本周开讲缠中说禅中医学</h2>
        # 提取标题
        pattern_time =re.compile('<title>(.*?)</title>',re.S)
        title = ''.join(re.findall(pattern_time,text))
        return title.split('_')[0]
    def getContent(self, page):
        '''
        将所有获得的文章内容整理在一起
        '''
        # 这是目录页里的标题内容，有些标题不全
        # pattern_title = re.compile(
        #     '<span class="atc_title">.*?<a.*?href.*?.html">(.*?)</a>.*?</span>', re.S)
        # all_titles = re.findall(pattern_title, page)
        
        pattern_url = re.compile(
            '<span class="atc_title">.*?<a.*?href="(.*?)">.*?</a>.*?</span>', re.S)
        all_urls = re.findall(pattern_url, page)
        # assert len(all_titles) == len(all_urls)
        all_texts = [self.getText(x) for x in all_urls]
        all_times = [self.getTime(x) for x in all_urls]
        all_titles = [self.getTitle(x) for x in all_urls]
        results = []
        for each_title, each_url, each_text, each_time in zip(all_titles, all_urls, all_texts, all_times):
            result = {
                'title': self.tool.replace(each_title),
                'url': each_url,
                'time': each_time,
                'text': each_text
            }
            results.append(result)
        return results


    def start(self):
        '''
        抓取文章
        '''
        indexPage = self.getPage(1)
        pageNum = self.getPageNum(indexPage)
        if pageNum == None:
            print ("网页链接URL已失效，请重试")
            return
        print ("该博客共有" + str(pageNum) + "页")
        self.all_results = []
        try:
            for i in range(1,int(pageNum)+1):
                print ("正在抓取第" + str(i) + "页数据")
                page = self.getPage(i)
                self.all_results += self.getContent(page)
        except IOError as e:
            print("抓取异常，原因为:" + e.message)
        finally:
            with open("./{}.json".format(self.fileName), 'w', encoding='utf-8') as json_file:
                json.dump(self.all_results, json_file, ensure_ascii=False, indent=4)
            print ("抓取任务完成,一共抓取了{}篇文章".format(len(self.all_results)))


    def write_word_file(self):  
        # 创建文档对象
        document = Document()
        # 全局设置字体为宋体
        document.styles['Normal'].font.name = '宋体'
        # 字号小四，对应12号字
        document.styles['Normal'].font.size = Pt(12)
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
        while self.all_results:
            each_result = self.all_results.pop()
            # 标题格式是标题2
            title_ = document.add_heading(level=2)
            # 添加标题内容
            title_run = title_.add_run(each_result['title'])  
            # 设置标题中文字体
            title_run.font.name = '宋体'  
            # 设置标题中文字体
            title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  
            # 字体颜色黑色
            title_run.font.color.rgb = RGBColor(0, 0, 0)
            p = document.add_paragraph(each_result['url'])
            p = document.add_paragraph(each_result['time'])
            for each_text in each_result['text']:
                p = document.add_paragraph(each_text)
                paragraph_format = p.paragraph_format
                # 首行缩进
                paragraph_format.first_line_indent = Inches(0.25)
                # 1.5倍行间距
                paragraph_format.line_spacing = 1.5 
            # 分页符
            document.add_page_break()
        document.save(self.fileName+'.docx')




 
if __name__ == '__main__':
    #### 文史哲学
    crawler = XLBKCrawler(
        # 博客目录链接
        baseUrl = 'http://blog.sina.com.cn/s/articlelist_1215172700_8_',
        # 从第一页开始爬取
        articleTag = 1,
        # 保存的文件名称
        fileName = '新浪博客文史哲学'
        )
    crawler.start()
    crawler.write_word_file()
    #### 时政经济
    crawler = XLBKCrawler(
        # 博客目录链接
        baseUrl = 'http://blog.sina.com.cn/s/articlelist_1215172700_10_',
        # 从第一页开始爬取
        articleTag = 1,
        # 保存的文件名称
        fileName = '新浪博客时政经济'
        )
    crawler.start()
    crawler.write_word_file()

    #### 音乐艺术
    crawler = XLBKCrawler(
        # 博客目录链接
        baseUrl = 'http://blog.sina.com.cn/s/articlelist_1215172700_7_',
        # 从第一页开始爬取
        articleTag = 1,
        # 保存的文件名称
        fileName = '新浪博客音乐艺术'
        )
    crawler.start()
    crawler.write_word_file()